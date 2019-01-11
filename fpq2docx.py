import sys, os
def resource_path(relative_path):
    """ Get absolute path to resource, works for dev and for PyInstaller """
    #try:
        # PyInstaller creates a temp folder and stores path in _MEIPASS
    #base_path = sys._MEIPASS
    #except Exception:
    base_path = os.path.abspath(".")

    return os.path.join(base_path, relative_path)

#print(resource_path('template.docx'))
import sys
import docx
from docx.shared import Cm
import yaml
from fpq import *

in_qfile = sys.argv[1]
in_afile = 'respostas.yaml'
outfile = in_qfile.split('.')[0]+".docx"


form = FPQ(in_qfile, in_afile)
estrutura = form.export()

def formatar_resposta(t,s):
    if t == "likert":
        return '\n'.join(['( ) {}'.format(c) for c in s])
    elif t == "relevância" :
        na = "( ) {}\n\n".format(s[0])
        opcoes = '\n'.join(["( ) {}".format(c) for c in s[1:]])
        return na + opcoes
    elif t == "escolha única":
        c0 = "Escolha uma das opções abaixo:\n\n"
        c1 = '\n'.join(["( ) {}".format(c) for c in s])
        return c0 + c1
    elif t == "escolha múltipla":
            c0 = "Marque as opções aplicáveis:\n\n"
            c1 = '\n'.join(["( ) {}".format(c) for c in s])
            return c0 + c1
    elif t == "composição":
        c0 = "(Os valores devem somar 100%)\n\n"
        m = lambda c: ' ' * (65 - len(c))
        c1 = '\n'.join(["{}:{} \t__%".format(c,m(c)) for c in s])
        return c0 + c1
    elif t == "valor":
        m = lambda c: ' ' * (65 - len(c))
        if len(s) == 1 and len(s[0]) == 0:
            c1 = "\t___________"
        else:
            c1 = '\n'.join(["{}:{} \t___________".format(c,m(c)) for c in s])
        return c1
    elif t == "talvez valor":
        m = lambda c: ' ' * (65 - len(c))
        c1 = '\n'.join(["( ) {}".format(c)for c in s[:-1]])

        m = lambda c: ' ' * (65 - len(c))
        c2 = "\n( ) {} - Informar valor: ___________".format(s[-1])
        return c1 + c2
    elif t == "lista":
        return '\n'.join(["-- {}".format(c) for c in s])
    else:
        c1 = '\n'.join(["( ) {}".format(c) for c in s])
        return c1

def inserir_resposta(document, tipo, resposta):
    out = formatar_resposta(tipo, resposta)
    par = document.add_paragraph(out)
    par.paragraph_format.left_indent = Cm(1)
    return par

with open(resource_path('template.docx'),'rb') as template:
    document = docx.Document(docx = template)
    #print(form.meta['título'])
    document.add_heading(form.meta['título'], 0)
    secoes = form.qstruct.ss

    for snum, s in enumerate(secoes):
        secao = secoes[s]

        heading = '{} - {}'.format(snum+1,secao['título'])
        document.add_heading(heading,1)
        for qnum, q in enumerate(estrutura[s]):
            questao = estrutura[s][q]
            heading = '{}.{}) {}'.format(snum+1,qnum+1,questao['texto'])
            document.add_heading(heading,2)
            par = inserir_resposta(document, questao['tipo-resposta'], questao['resposta'])
    document.save(form.meta['título']+".docx")
